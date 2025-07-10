!pip install fpdf
!pip install python-docx
!pip install PyPDF2

import requests #for perplexity
from fpdf import FPDF #to download sop as pdf
from docx import Document #to download sop as doc
import re 
import os
import PyPDF2 #for cv upload and parsing

#to extract name from cv
def extract_name_from_cv(text):
    lines = text.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]

    #strategy 1: Look for title-style capitalized name line on top
    for line in lines[:5]:
        if (
            re.match(r"^[A-Z][a-z]+(?: [A-Z][a-z]+)+$", line) and
            len(line.split()) <= 4 and
            not any(char in line for char in ['|', '@', 'http', '/', '\\', ':'])
        ):
            return line

    #strategy 2: Try extracting using a known pattern (like "Name: John Doe")
    match = re.search(r'Name[:\-]\s*(.+)', text, re.IGNORECASE)
    if match:
        possible_name = match.group(1).strip()
        if len(possible_name.split()) <= 4:
            return possible_name

    return "Your Name"

#to extract academic qualifications
def extract_academic_qualifications(text):
    match = re.search(r'(EDUCATION|ACADEMIC QUALIFICATIONS)(.*?)(PROJECTS|SKILLS|EXPERIENCE|ACHIEVEMENTS|$)', text, re.IGNORECASE | re.DOTALL)
    if match:
        degrees = match.group(2).strip()
        degrees = re.sub(r'\n+', '\n', degrees)
        return degrees
    return ""

#to determine intended degreee
def determine_intended_degree(academic_text):
    academic_text = academic_text.lower()
    if "phd" in academic_text:
        return "PhD"
    elif "master" in academic_text or "msc" in academic_text:
        return "PhD"
    elif "bachelor" in academic_text:
        return "Masters"
    else:
        return "Masters"

#function to extract text from document or pdf
def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    except:
        return ""

def extract_text_from_pdf(filepath):
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    except:
        return ""

#parse details from extracted text
def extract_section(text, headings):
    pattern = '|'.join(headings)
    match = re.search(rf'(?i)({pattern})[:\s]*([\s\S]+?)(?:\n\s*\n|$)', text)
    return match.group(2).strip() if match else None
    
def parse_cv(text):
    return {
        "name": extract_name_from_cv(text),
        "academic_qualifications": extract_academic_qualifications(text),
        "intended_degree": determine_intended_degree(extract_academic_qualifications(text)),
        "key_skills": extract_section(text, ["skills", "technical skills"]),
        "projects": extract_section(text, ["projects", "publications", "research"]),
        "awards": extract_section(text, ["awards", "scholarships", "recognitions"]),
        "hobbies": extract_section(text, ["hobbies", "volunteer work", "extracurriculars"])
    }

#function to print ask input (either parsed or user type)
def ask_input(field, default=None, optional=False, options=None):
    while True:
        if default:
            prompt = f"Enter {field} (Press Enter to keep: '{default}')"
        elif options:
            prompt = f"Select {field} ({'/'.join(options)})"
        else:
            prompt = f"Enter {field}{' (optional)' if optional else ''}:"
        print(prompt)
        val = input().strip()
        if val:
            return val
        elif default is not None:
            return default
        elif optional:
            return ""
        else:
            print(f"{field} is required. Please enter a value.")

#function to get user inputs
def collect_user_inputs():
    print("Do you want to (1) Upload CV or (2) Enter details manually?")
    method = input("Enter 1 or 2: ").strip()

    parsed_data = {}
    if method == '1':
        filepath = input("Enter path to your CV (pdf or docx): ").strip()
        if not os.path.exists(filepath):
            print("File not found. Switching to manual input.")
        else:
            if filepath.endswith(".pdf"):
                text = extract_text_from_pdf(filepath)
            elif filepath.endswith(".docx"):
                text = extract_text_from_docx(filepath)
            else:
                print("Unsupported format. Switching to manual input.")
                text = ""
            parsed_data = parse_cv(text)

    inputs = {}
    inputs["word_count_target"] = ask_input("Word Count Target", optional=True)
    inputs["tone"] = ask_input("Tone (optional)", optional=True, options=["Formal", "Creative", "Concise", "Detailed", "Custom"])

    inputs["name"] = ask_input("Name", default=parsed_data.get("name"))
    inputs["country_of_origin"] = ask_input("Country of origin")
    inputs["intended_degree"] = ask_input("Intended Degree", default=parsed_data.get("intended_degree"))
    inputs["preferred_country"] = ask_input("Preferred Country of Study")
    inputs["field_of_study"] = ask_input("Preferred Field of Study")
    inputs["preferred_uni"] = ask_input("Preferred University")

    inputs["academic_qualifications"] = ask_input("Academic qualifications (Degree, uni, year, subjects)", default=parsed_data.get("academic_qualifications"))
    inputs["key_skills"] = ask_input("Key skills", default=parsed_data.get("key_skills"))
    inputs["strengths"] = ask_input("Strengths", optional=True)
    inputs["why_field"] = ask_input("Why this field?", optional=True)
    inputs["why_uni"] = ask_input("Why this university?", optional=True)
    inputs["projects"] = ask_input("Projects / research / publications", default=parsed_data.get("projects"), optional=True)
    inputs["awards"] = ask_input("Awards / scholarships / recognitions", default=parsed_data.get("awards"), optional=True)
    inputs["goals"] = ask_input("Long term goals", optional=True)
    inputs["hobbies"] = ask_input("Hobbies / volunteer work / extracurriculars", default=parsed_data.get("hobbies"), optional=True)
    inputs["challenge"] = ask_input("A time you overcame a challenge", optional=True)

    return inputs
def build_sop_prompt(user_inputs):
    name = user_inputs.get("name")
    country_of_origin = user_inputs.get("country_of_origin")
    intended_degree = user_inputs.get("intended_degree")
    preferred_country = user_inputs.get("preferred_country")
    word_count = user_inputs.get("word_count_target")
    tone = user_inputs.get("tone") or "Formal"
    field_of_study = user_inputs.get("field_of_study")
    preferred_uni = user_inputs.get("preferred_uni")

    base_prompt = (
        f"I am {name}, I am from {country_of_origin}. I want to study {intended_degree} in {preferred_country}. "
        f"My preferred field of study is {field_of_study}. "
        f"My preferred university is {preferred_uni}. "
        f"I want you to write me a SOP"
    )

    if word_count:
        base_prompt += f" following word limit count {word_count}."

    base_prompt += (
        "Make sure the SOP is ATS friendly and does not look like an AI wrote it. "
        "It should look like a human wrote it. "
        f"The SOP should follow this tone {tone}. "
        "Only respond with the SOP text. Do not include any explanations or additional messages."
        "Exclude all inline citations or footnote markers.\n"
    )

    base_prompt += "Here are my details:\n"

    optional_fields = [
        ("academic_qualifications", "My academic qualifications are"),
        ("key_skills", "My key skills are"),
        ("strengths", "My strengths are"),
        ("why_field", "I want to pursue this field because"),
        ("why_uni", "And this university because"),
        ("projects", "I have done"),
        ("awards", "I have received"),
        ("goals", "My long term goals are"),
        ("hobbies", "I like to"),
        ("challenge", "More about me:")
    ]

    for key, label in optional_fields:
        value = user_inputs.get(key)
        if value:
            base_prompt += f"{label} {value}.\n"

    return base_prompt.strip()

def call_perplexity_api(prompt, token):
    url = "https://api.perplexity.ai/chat/completions"
    payload = {
        "model": "sonar",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2048
    }
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    response = requests.post(url, json=payload, headers=headers)
    return response.json()

def clean_text_for_pdf(text):
    text = text.replace("—", "-").replace("–", "-")
    text = text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    return text.encode("latin-1", "ignore").decode("latin-1")

def save_pdf(filename, content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    content = clean_text_for_pdf(content)
    pdf.multi_cell(0, 10, content)
    pdf.output(filename)

def save_docx(filename, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)

def main():
    user_inputs = collect_user_inputs()
    token = "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"
    prompt = build_sop_prompt(user_inputs)
    response = call_perplexity_api(prompt, token)
    sop = response.get("choices", [{}])[0].get("message", {}).get("content", "").strip()

    if not sop:
        print("No SOP generated. Please try again.")
        return

    word_count = len(re.findall(r'\w+', sop))
    print("\n--- Generated SOP ---\n")
    print(sop)
    print(f"\n--- Word count: {word_count} ---\n")

    while True:
        print("Do you want to (1) See the SOP again, (2) Download the SOP, or (3) Exit?")
        choice = input("Enter 1, 2, or 3: ").strip()

        if choice == '1':
            print("\n--- Generated SOP ---\n")
            print(sop)
            print(f"\n--- Word count: {word_count} ---\n")
        elif choice == '2':
            print("Download as (pdf/doc)? Enter 'pdf' or 'doc':")
            file_type = input().strip().lower()
            if file_type == 'pdf':
                filename = "SOP.pdf"
                save_pdf(filename, sop)
                print(f"PDF saved as {filename}")
            elif file_type == 'doc':
                filename = "SOP.docx"
                save_docx(filename, sop)
                print(f"DOC saved as {filename}")
            else:
                print("Invalid file type selected.")
        elif choice == '3':
            print("Exiting...")
            break
        else:
            print("Invalid choice, please enter 1, 2, or 3.")

if __name__ == "__main__":
    main()
